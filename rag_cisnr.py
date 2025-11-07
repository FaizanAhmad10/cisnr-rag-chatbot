import os
from dotenv import load_dotenv
from docx import Document
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

# -------------------- LOAD API KEY --------------------
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

# -------------------- EXTRACT TEXT FROM WORD FILE --------------------
def extract_text_from_docx(file_path):
    """Extract all text from a Word document."""
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Also extract text from tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n\n'.join(full_text)
    
    except FileNotFoundError:
        print(f"❌ Error: File '{file_path}' not found!")
        print(f"   Please make sure the file exists in: {os.path.abspath(file_path)}")
        exit(1)
    except Exception as e:
        print(f"❌ Error reading Word file: {e}")
        exit(1)

# -------------------- LOAD DATA FROM WORD FILE --------------------
print("📄 Loading data from Word file...")
word_file = 'cisnr_data.docx'  # Change this to your file name if different
cisnr_data = extract_text_from_docx(word_file)
print(f"✅ Loaded {len(cisnr_data)} characters from Word file")
print(f"   Preview: {cisnr_data[:100]}...\n")

# -------------------- TEXT SPLITTING --------------------
print("🔄 Splitting text into chunks...")
splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
docs = splitter.split_documents([LangchainDocument(page_content=cisnr_data)])
print(f"✅ Created {len(docs)} text chunks\n")

# -------------------- EMBEDDINGS --------------------
print("🔄 Creating embeddings...")
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# -------------------- VECTOR STORE --------------------
print("🔄 Building vector store...")
persist_directory = "vector_store"
vectorstore = Chroma.from_documents(
    documents=docs,
    embedding=embeddings,
    persist_directory=persist_directory
)
print("✅ Vector store created successfully\n")

# -------------------- GROQ LLM --------------------
llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0, api_key=groq_api_key)

# -------------------- RAG CHAIN --------------------
retriever = vectorstore.as_retriever(search_kwargs={"k": 3})

# Create prompt template
prompt = ChatPromptTemplate.from_template("""
Answer the question based only on the following context:

{context}

Question: {question}

Answer:
""")

# Helper function to format documents
def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

# Build RAG chain using LCEL
rag_chain = (
    {"context": retriever | format_docs, "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()
)

# -------------------- CHAT LOOP --------------------
print("=" * 60)
print("🤖 CISNR RAG Chatbot (Groq + ChromaDB)")
print("=" * 60)
print("Type 'exit' to quit.\n")

while True:
    query = input("You: ")
    if query.lower() in ["exit", "quit"]:
        print("\n👋 Goodbye!")
        break
    try:
        answer = rag_chain.invoke(query)
        print(f"Bot: {answer}\n")
    except Exception as e:
        print(f"❌ Error: {e}\n")



